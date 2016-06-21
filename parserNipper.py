#!/usr/bin/python
#!/usr/bin/env python

"""parserNipper is a script to parse Nipper XML files"""

from xml.etree import ElementTree
import argparse
import docx
import os
import logging
import sys
import readline

# Requires python-docx library, apt-get update; apt-get install -y python-pip;pip install python-docx

#################################################
#                    Variables                  #
#################################################
__author__ = "Russel Van Tuyl"
__license__ = "GPL"
__version__ = "0.1"
__maintainer__ = "Russel Van Tuyl"
__email__ = "Russel.VanTuyl@gmail.com"
__status__ = "Development"
logging.basicConfig(stream=sys.stdout, format='%(asctime)s\t%(levelname)s\t%(message)s',
                    datefmt='%Y-%m-%d %I:%M:%S %p', level=logging.WARNING)  # Log to STDOUT

nipper_xml = None
readline.parse_and_bind('tab: complete')
readline.set_completer_delims('\t')
DEBUG = False
VERBOSE = False
#################################################
#                   COLORS                      #
#################################################
note = "[\033[0;0;33m-\033[0m]"
warn = "[\033[0;0;31m!\033[0m]"
info = "[\033[0;0;36mi\033[0m]"
question = "[\033[0;0;37m?\033[0m]"


def parse_xml(xml_file):
    """Parse XML file and return ElementTree object"""
    logging.info(str(xml_file))
    global nipper_xml
    xml_tree = ElementTree.parse(xml_file)

    nipper_xml = xml_tree.getroot()


def get_devices_information():
    """Create a dictionary of devices found in the Nipper XML from the /document/information/devices section"""
    global nipper_xml
    devices = {}

    for device in nipper_xml.findall('./information/devices/device'):
        if DEBUG:
            print "\t" + note + "Name: %s" % device.get('name')
            print "\t" + note + "Type: %s" % device.get('type')
            print "\t" + note + "OS: %s" % device.get('os')
            print "\t" + note + "OS Version: %s" % device.get('osversion')
        devices[device.attrib.get('name')] = {'name': device.get('name'),
                                                         'type': device.get('type'),
                                                         'os': device.get('os'),
                                                         'osversion': device.get('osversion')}
    if DEBUG:
        print info + "Device Object:"
        print devices
        raw_input(warn + "Press enter to continue")
    return devices


def get_devices_summary():
    """Create a dictionary of devices found in the Nipper XML from the /summary"""

    # This function was created to replace get_devices_information
    # because it wasn't detecting virtual systems in Palo Alto Virtual Systems
    global nipper_xml
    devices = {}
    headings = []

    # Add the table headings to a list
    for h in nipper_xml.findall("./summary/table/[@ref='SCOPE.AUDITDEVICELIST.TABLE']/headings/heading"):
        if h not in headings:
            headings.append(h.text)

    for device in nipper_xml.findall("./summary/table/[@ref='SCOPE.AUDITDEVICELIST.TABLE']/tablebody/tablerow"):
        values = []
        for i in device.findall('./tablecell/item'):
            if i not in values:
                values.append(i.text)
        if DEBUG:
            print "\t" + note + "%s: %s" % (headings[headings.index('Name')], values[headings.index('Name')])
            print "\t" + note + "%s: %s" % (headings[headings.index('Device')], values[headings.index('Device')])
            print "\t" + note + "%s: %s" % (headings[headings.index('OS')], values[headings.index('OS')].split(" ")[0])
            print "\t" + note + "%s: %s" % (headings[headings.index('OS')], values[headings.index('OS')].split(" ")[1])
        devices[values[headings.index('Name')]] = {'name': values[headings.index('Name')],
                                                   'type': values[headings.index('Device')],
                                                   'os': values[headings.index('OS')].split(' ')[0],
                                                   'osversion': values[headings.index('OS')].split(' ')[1]
                                                   }

    if DEBUG:
        print info + "Device Object:"
        print devices
        raw_input(warn + "Press enter to continue")
    return devices


def get_ipv4_interfaces(device_name):
    """Build a list of interfaces per device"""
    interfaces = {}
    if DEBUG:
        print note + "Entering into get_ipv4_interfaces function"
    # Needs to be fixed. Get list of interfaces first, then IP addresses, then VLAN, then ACLs
    config_element = nipper_xml.find("./report/part/[@ref='CONFIGURATION']")

    for section in config_element.findall('./section'):
        device_item = None

        for i in section.get('title').split():
            if device_name == i:
                device_item = device_name
                if DEBUG:
                    print "\t" + note + "Set Device: %s" % device_name

        if device_item is not None:
            interface_element = section.find("./section/[@ref='CONFIGURATION.ADDRESSES']/section/"
                                             "[@ref='ADDRESSES.IPV4']")
            if interface_element is not None:
                headings = []
                items = []
                for heading in interface_element.findall("./table/[@title='IPv4 addresses']/headings/heading"):
                    headings.append(heading.text)
                    if DEBUG:
                        print "\t" + note + "Set Heading: %s" % heading.text
                for item in interface_element.findall("./table/[@title='IPv4 addresses']/tablebody"
                                                      "/tablerow/tablecell/item"):
                    items.append(item.text)
                    if DEBUG:
                        print "\t" + note + "Set Item: %s" % item.text
                i = 0
                interface_id = None
                if DEBUG:
                    print "\t" + note + "Heading List: %s" % headings
                    print "\t" + note + "Items List: %s" % items
                for item in items:
                    if i > (len(headings) - 1):
                        i = 0
                    if DEBUG:
                        print "\t" + info + "Heading: %s\t Item: %s" % (headings[i], item)
                    if i is 0:
                        interface_id = item
                        interfaces[interface_id] = {}
                    interfaces[interface_id].update({headings[i]: item})
                    i += 1

            interfaces_element = section.find("./section/[@ref='CONFIGURATION.INTERFACES']/section/"
                                              "[@ref='ETHINTERFACESLAYER3']")
            if interfaces_element is not None:
                headings = []
                for heading in interfaces_element.findall("./table/[@title='Layer 3 Ethernet Interfaces']"
                                                          "/headings/heading"):
                    headings.append(heading.text)
                for tr in interfaces_element.findall("./table/[@title='Layer 3 Ethernet Interfaces']"
                                                     "/tablebody/tablerow"):
                    items = []
                    for i in tr.findall("./tablecell/item"):
                        items.append(i.text)
                    if 'Zone' in headings:
                        interfaces[items[headings.index('Interface')]].update({'Zone': items[headings.index('Zone')]})
                    if 'VLAN' in headings:
                        interfaces[items[headings.index('Interface')]].update({'VLAN': items[headings.index('VLAN')]})
    if DEBUG:
        print info + "Interfaces object: "
        print interfaces
        raw_input(warn + "Press any key to continue")
    return interfaces


def get_cve(device_list):
    """Get a list of all CVEs"""

    # A list of devices has to be passed in to match device names in strings. Nipper does not do a good job of making a
    # unique identifier for an affected host to a CVE. They just write a sentence with the device name in it.

    cves = {}

    cve_element = nipper_xml.find("./report/part/[@ref='VULNAUDIT']")

    for section in cve_element.findall('./section'):
        # print section.get('title')
        if section.get('title').startswith('CVE-'):  # Look only at CVE sections
            cve = section.get('title')
            cves[cve] = {}
            if DEBUG:
                print info + "CVE: %s" % cve
            # CVSS_v2 Score
            cvss_score = section.find("./infobox/infodata/[@label='CVSSv2 Score']").text
            cves[cve].update({'CVSSv2_Score': cvss_score})
            if DEBUG:
                print "\t" + info + "CVSSv2 Score: %s" % cvss_score
            # Single Devices
            for ad_section in section.findall("./section/[@title='Affected Device']"):
                for d in device_list:
                    for i in ad_section.find("./text").text.split():
                        if d == i:
                            cves[cve].update({'Hosts': [d]})
                            if DEBUG:
                                print "\t" + note + "Single Device: %s" % d
            # Multiple Devices
            ad_list = []
            for ad_section in section.findall("./section/[@title='Affected Devices']/list/listitem"):
                ad_list.append(ad_section.text.split(" - ")[1].rstrip(";").rstrip("."))
            if len(ad_list) > 0:
                cves[cve].update({'Hosts': ad_list})
                if DEBUG:
                    print "\t" + note + "Multiple Devices: %s" % ad_list
    if DEBUG:
        print info + "CVE Object:"
        print cves
        raw_input(warn + "Press enter to continue")
    return cves


def get_nipper_version():
    """Get the version of Nipper used to generate the XML file"""
    nipper_version = nipper_xml.find("./information/generator/version").text
    return nipper_version

# TODO Get rid of the get_any_source_any_destination_any_port function; get-filter_rules is used instead
def get_any_source_any_destination_any_port(devices):
    """Get all firewall rules with ANY Source, ANY Destination, and ANY Port for the passed in devices"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='FILTER.RULE.AEAA']/"
                                "section[@ref='FINDING']/table")

    rules = {}

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                acl = table.get('ref').lstrip('FILTER.RULE.AEAA')[:-1]
                rules[acl] = {}
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                if DEBUG:
                    print info + "%s ACL: %s" % (device, table.get('ref').lstrip('FILTER.RULE.AEAA')[:-1])
                    print info + "HEADINGS: %s" % headings
                for row in table.findall('./tablebody/tablerow'):
                    i = 0
                    rule_id = None
                    for data in row.findall('./tablecell'):
                        print "DATA: %s" % data.find('./item').text
                        if headings[i] == 'Rule':
                            rule_id = data.find('./item').text
                            rules[acl][rule_id] = {'Device': device}
                        if rule_id:
                            rules[acl][rule_id].update({headings[i]: data.find('./item').text})
                        i += 1
    if DEBUG:
        print info + "Any Source/Any Destination/Any Port Rule Object:"
        print rules
        raw_input(warn + "Press enter to continue")
    return rules

# TODO Get rid of the get_any_source_fixed_destination_any_port function
def get_any_source_fixed_destination_any_port(devices):
    """Get all firewall rules with ANY Source, a FIXED Destination, and ANY Port for the passed in devices"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='FILTER.RULE.AENA']/"
                                "section[@ref='FINDING']/table")

    rules = {}

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                acl = table.get('ref').lstrip('FILTER.RULE.AENA')[:-1]
                rules[acl] = {}
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                if DEBUG:
                    print info + "%s ACL: %s" % (device, table.get('ref').lstrip('FILTER.RULE.AENA')[:-1])
                    print info + "HEADINGS: %s" % headings
                for row in table.findall('./tablebody/tablerow'):
                    i = 0
                    rule_id = None
                    for data in row.findall('./tablecell'):
                        print "DATA: %s" % data.find('./item').text
                        if headings[i] == 'Rule':
                            rule_id = data.find('./item').text
                            rules[acl][rule_id] = {'Device': device}
                        if rule_id:
                            rules[acl][rule_id].update({headings[i]: data.find('./item').text})
                        i += 1
    if DEBUG:
        print info + "Any Source/Fixed Destination/Any Port Rule Object:"
        print rules
        raw_input(warn + "Press enter to continue")
    return rules

# TODO Get rid of the get_fixed_source_any_destination_any_port function
def get_fixed_source_any_destination_any_port(devices):
    """Get all firewall rules with a fixed Source, ANY Destination, and ANY Port for the passed in devices"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='FILTER.RULE.NEAA']/"
                                "section[@ref='FINDING']/table")

    rules = {}

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                acl = table.get('ref').lstrip('FILTER.RULE.NEAA')[:-1]
                rules[acl] = {}
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                if DEBUG:
                    print info + "%s ACL: %s" % (device, table.get('ref').lstrip('FILTER.RULE.NEAA')[:-1])
                    print info + "HEADINGS: %s" % headings
                for row in table.findall('./tablebody/tablerow'):
                    i = 0
                    rule_id = None
                    for data in row.findall('./tablecell'):
                        print "DATA: %s" % data.find('./item').text
                        if headings[i] == 'Rule':
                            rule_id = data.find('./item').text
                            rules[acl][rule_id] = {'Device': device}
                        if rule_id:
                            rules[acl][rule_id].update({headings[i]: data.find('./item').text})
                        i += 1
    if DEBUG:
        print info + "Fixed Source, Any Destination, Any Port Rule Object:"
        print rules
        raw_input(warn + "Press enter to continue")
    return rules

# TODO Get rid of the get_fixed_source_fixed_destination_any_port funciton
def get_fixed_source_fixed_destination_any_port(devices):
    """Get all firewall rules with a fixed Source, fixed Destination, and ANY Port for the passed in devices"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='FILTER.RULE.NENA']/"
                                "section[@ref='FINDING']/table")

    rules = {}

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                acl = table.get('ref').lstrip('FILTER.RULE.NENA')[:-1]
                rules[acl] = {}
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                if DEBUG:
                    print info + "%s ACL: %s" % (device, table.get('ref').lstrip('FILTER.RULE.NENA')[:-1])
                    print info + "HEADINGS: %s" % headings
                for row in table.findall('./tablebody/tablerow'):
                    i = 0
                    rule_id = None
                    for data in row.findall('./tablecell'):
                        print "DATA: %s" % data.find('./item').text
                        if headings[i] == 'Rule':
                            rule_id = data.find('./item').text
                            rules[acl][rule_id] = {'Device': device}
                        if rule_id:
                            rules[acl][rule_id].update({headings[i]: data.find('./item').text})
                        i += 1
    if DEBUG:
        print info + "Fixed Source, Fixed Destination, Any Port Rule Object:"
        print rules
        raw_input(warn + "Press enter to continue")
    return rules


def get_filter_rules(devices, filter_name):
    """Get all firewall rules for the passed in devices and filter name"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='" + filter_name + "']/"
                                "section[@ref='FINDING']/table")

    rules = {}

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                acl = table.get('ref').lstrip(filter_name)[:-1]
                rules[acl] = {}
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                if DEBUG:
                    print info + "%s ACL: %s" % (device, table.get('ref').lstrip(filter_name)[:-1])
                    print info + "HEADINGS: %s" % headings
                for row in table.findall('./tablebody/tablerow'):
                    i = 0
                    rule_id = None
                    for data in row.findall('./tablecell'):
                        if headings[i] == 'Rule':
                            rule_id = data.find('./item').text
                            rules[acl][rule_id] = {'Device': device}
                        if rule_id:
                            rules[acl][rule_id].update({headings[i]: data.find('./item').text})
                        i += 1
    if DEBUG:
        print info + "%s Rule Object: " % filter_name
        print rules
        raw_input(warn + "Press enter to continue")
    return rules


def get_vpn_aggressive(devices):
    """Get VPN Agressive mode vulnerability"""

    section = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='REMOTEACCESS.AGGRESSIVEMODE']"
                                 "/section/[@title='Affected Devices']/list/listitem")

    vulns = []
    if section:
        for device in devices:
            for item in section:
                if device in item.text.split():
                    vulns.append(device)

    if DEBUG:
        print note + "VPN Agrressive Mode Object: ", vulns
        raw_input(warn + "Press enter to continue...")
    if len(vulns) > 0:
        return vulns
    else:
        return None


def get_weak_snmp_community(devices):
    """Get Weak SNMP Community String informaiton from report"""

    tables = nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section/[@ref='SNMP.WEAK.COMMUNITY']/"
                                "section[@ref='FINDING']/table")

    snmp = {}

    # Data Structure
    # <device ID> : {}
    #   <Community String>: {}
    #       <Weakness>: weakness
    #       <Host>: host

    for table in tables:
        for device in devices:
            if device in table.get('title').split():
                headings = []
                data = []
                for heading in table.findall('./headings/heading'):
                    headings.append(heading.text)
                for row in table.findall('./tablebody/tablerow'):
                    for item in row.findall('.tablecell'):
                        data.append(item.find('./item').text)
                if DEBUG:
                    print info + "SNMP Weak Community String:"
                    print "\t" + info + "Headers: %s" % headings
                    print "\t" + info + "Data: %s" % data
                    raw_input(warn + "Press enter to continue...")
                if device not in snmp:
                    snmp[device] = {}
                c = headings.index('Community')
                w = headings.index('Weakness')
                snmp[device][data[c]] = {headings[w]: data[w]}
                if 'Host' in headings:
                    h = headings.index('Host')
                    snmp[device][data[c]].update({headings[h]: data[h]})

    if DEBUG:
        print "SNMP Weak Community String Dict:"
        print info, snmp
        raw_input(warn + "Press enter to continue...")

    return snmp


def write_intro(report, devices):
    """Write the introduction paragraph for a firewall configuration audit narrative"""

    report.add_heading("Firewall Configuration Audit")

    report.add_paragraph("For the firewall configuration audit, the assessment team reviewed the configuration of "
                         "%d firewall devices. The table below shows general device information for "
                         "each of the assessed devices." % len(devices), style='Normal')

    table = report.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Version'
    table.style = 'Medium Grid 1 Accent 1'

    for device in devices:
        row_cells = table.add_row().cells
        row_cells[0].text = "%s" % devices[device]['name']
        row_cells[1].text = "%s" % devices[device]['type']
        row_cells[2].text = "%s" % devices[device]['osversion']

    report.add_paragraph("\nThe assessment team leveraged extensive experience as firewall administrators as well as "
                         "commercial tools to perform the firewall device audits and provide security recommendations. "
                         "Output from the commercial tools has been provided as supplemental documentation "
                         "to this report.", style='Normal')

    return report


def write_interfaces(report, devices):
    """Write the Active Interfaces portion of the report"""

    report.add_heading("Device Active Interfaces", 3)

    report.add_paragraph("The following device information shows each active interface configuration as well as the "
                         "corresponding zone defined for that interface. During the configuration audit, the assessment"
                         " team placed emphasis on the active interfaces identified in the following tables.",
                         style='Normal')

    for device in devices:
        if len(devices[device]['Interfaces']) > 0:  # Don't create a table when there are no interfaces
            report.add_heading("\n%s Active Interfaces" % device, 4)
            table = report.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Interface'
            hdr_cells[1].text = 'Address'
            table.style = 'Medium Grid 1 Accent 1'
            for interface in devices[device]['Interfaces']:
                if 'Active' in devices[device]['Interfaces'][interface].keys():
                    if devices[device]['Interfaces'][interface]['Active']:
                        if devices[device]['Interfaces'][interface]['Active'].upper() == "YES":
                            row_cells = table.add_row().cells
                            row_cells[0].text = "%s" % devices[device]['Interfaces'][interface]['Interface']
                            if 'Address' in devices[device]['Interfaces'][interface].keys():
                                row_cells[1].text = "%s" % devices[device]['Interfaces'][interface]['Address']
                            if "ACL In" in devices[device]['Interfaces'][interface].keys():
                                hdr_cells[2].text = 'ACL IN'
                                if devices[device]['Interfaces'][interface]['ACL In'] is not None:
                                    row_cells[2].text = "%s" % devices[device]['Interfaces'][interface]['ACL In']
                                else:
                                    row_cells[2].text = "---"
                            elif "Zone" in devices[device]['Interfaces'][interface].keys():
                                hdr_cells[2].text = 'Zone'
                                if devices[device]['Interfaces'][interface]['Zone'] is not None:
                                    row_cells[2].text = "%s" % devices[device]['Interfaces'][interface]['Zone']
                                else:
                                    row_cells[2].text = "---"
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = "%s" % devices[device]['Interfaces'][interface]['Interface']
                    if 'Address' in devices[device]['Interfaces'][interface].keys():
                        row_cells[1].text = "%s" % devices[device]['Interfaces'][interface]['Address']
                    if "ACL In" in devices[device]['Interfaces'][interface].keys():
                        hdr_cells[2].text = 'ACL IN'
                        row_cells[2].text = "%s" % devices[device]['Interfaces'][interface]['ACL In']
                    else:
                        hdr_cells[2].text = 'DELETE ME'
                        row_cells[2].text = "---"
        else:
            logging.info("%s has no interfaces. An interfaces table will not be generated for this device" % device)
            if DEBUG:
                print warn + "%s has no interfaces" % device
    return report


def write_cve(report, cves):
    """Write the CVE section of the report"""

    # Build CVE Table for report
    report.add_heading("Software Vulnerabilities", 3)

    report.add_paragraph("The assessment team identified a cumulative total of %d documented common vulnerabilities and"
                         " exposures (CVEs) affecting the evaluated firewall devices as shown below in TABLE X. "
                         "Additionally, the corresponding common vulnerability scoring system version two (CVSSv2) "
                         "score is provided alongside the vulnerability identifier." % len(cves), style='Normal')

    table = report.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'CVSSv2 Score'
    hdr_cells[1].text = 'CVE ID'
    hdr_cells[2].text = 'Affected Host(s)'
    table.style = 'Medium Grid 1 Accent 1'
    cve_sorted_keys = sorted(cves, key=lambda x: cves[x]['CVSSv2_Score'], reverse=True)
    for key in cve_sorted_keys:
        row_cells = table.add_row().cells
        row_cells[0].text = str(cves[key]['CVSSv2_Score'])
        row_cells[1].text = key
        affected_hosts = ""
        if cves[key]['Hosts'] is not None:
            for host in cves[key]['Hosts']:
                affected_hosts += host
                if len(cves[key]['Hosts']) is 1 or cves[key]['Hosts'].index(host) is (len(cves[key]['Hosts']) - 1):
                    pass
                else:
                    affected_hosts += ", "
        row_cells[2].text = affected_hosts

    return report


def write_filter_rules(report, filter_rule_vulns):
    """Write filter rules section of the report"""

    report.add_heading("Overly Permissive Firewall Entries", 3)

    report.add_paragraph("The assessment team identified several access control lists (ACL) configured with overly "
                         "permissive rule entries on the assessed device(s). Overly permissive rules allow a "
                         "combination of traffic to or from ANY source IP, ANY destination IP, and/or ANY destination "
                         "port. Rules should be configured as restrictive as possible, allowing for precise control "
                         "of organizational traffic while facilitating business operations.  Rules that allow any "
                         "type of traffic, or traffic to any hosts, can often be more specific. Rules that allow "
                         "traffic from untrusted sources to trusted destinations should be extremely restrictive.  "
                         "Unrestrictive entries may facilitate unauthorized access to systems or enable attackers "
                         "to pivot through a network.", style='Normal')

    any_source_any_destination_table = None
    any_source_any_port_table = None
    any_destination_any_port_table = None
    any_port_table = None

    for vuln in filter_rule_vulns:
        if DEBUG:
            print filter_rule_vulns[vuln]
            print info + "write_filter_rules VULN: %s" % vuln
            raw_input(warn + "Press enter to continue")
        if (vuln == 'AEAA' or vuln == 'AEAE') and any_source_any_destination_table is None:
            report.add_heading("Any Source/Any Destination", 4)
            any_source_any_destination_table = report.add_table(rows=1, cols=3)
            hdr_cells = any_source_any_destination_table.rows[0].cells
            hdr_cells[0].text = 'Host'
            hdr_cells[1].text = 'ACL'
            hdr_cells[2].text = 'Rule(s)'
            any_source_any_destination_table.style = 'Medium Grid 1 Accent 1'
            if DEBUG:
                print info + "CREATED AEAA/AEAE TABLE"
        elif vuln == 'AENA' and any_source_any_port_table is None:
            report.add_heading("Any Source/Any Port", 4)
            any_source_any_port_table = report.add_table(rows=1, cols=3)
            hdr_cells = any_source_any_port_table.rows[0].cells
            hdr_cells[0].text = 'Host'
            hdr_cells[1].text = 'ACL'
            hdr_cells[2].text = 'Rule(s)'
            any_source_any_port_table.style = 'Medium Grid 1 Accent 1'
            if DEBUG:
                print info + "CREATED AENA TABLE"
        elif (vuln == 'NEAA' or vuln == 'EEAA') and any_destination_any_port_table is None:
            report.add_heading("Any Destination/Any Port", 4)
            any_destination_any_port_table = report.add_table(rows=1, cols=3)
            hdr_cells = any_destination_any_port_table.rows[0].cells
            hdr_cells[0].text = 'Host'
            hdr_cells[1].text = 'ACL'
            hdr_cells[2].text = 'Rule(s)'
            any_destination_any_port_table.style = 'Medium Grid 1 Accent 1'
            if DEBUG:
                print info + "CREATED NEAA/EEAA TABLE"
        elif vuln == 'NENA' and any_port_table is None:
            report.add_heading("Any Port", 4)
            any_port_table = report.add_table(rows=1, cols=3)
            hdr_cells = any_port_table.rows[0].cells
            hdr_cells[0].text = 'Host'
            hdr_cells[1].text = 'ACL'
            hdr_cells[2].text = 'Rule(s)'
            any_port_table.style = 'Medium Grid 1 Accent 1'
            if DEBUG:
                print info + "CREATED NENA TABLE"
        if (vuln == 'AEAA' or vuln == 'AEAE') and any_source_any_destination_table is not None:
            for acl in filter_rule_vulns[vuln]:
                row_cells = any_source_any_destination_table.add_row().cells
                row_cells[1].text = "%s" % acl
                rules = ""
                rule_device = None
                for rule in filter_rule_vulns[vuln][acl]:
                    if len(rules) > 0:
                        rules += ", %s" % rule
                    else:
                        rules += "%s" % rule
                    if rule_device is None:
                        if filter_rule_vulns[vuln][acl][rule]['Device']:
                            row_cells[0].text = "%s" % filter_rule_vulns[vuln][acl][rule]['Device']
                row_cells[2].text = "%s" % rules
        elif vuln == 'AENA' and any_source_any_port_table is not None:
            for acl in filter_rule_vulns[vuln]:
                row_cells = any_source_any_port_table.add_row().cells
                row_cells[1].text = "%s" % acl
                rules = ""
                rule_device = None
                for rule in filter_rule_vulns[vuln][acl]:
                    if len(rules) > 0:
                        rules += ", %s" % rule
                    else:
                        rules += "%s" % rule
                    if rule_device is None:
                        if filter_rule_vulns[vuln][acl][rule]['Device']:
                            row_cells[0].text = "%s" % filter_rule_vulns[vuln][acl][rule]['Device']
                row_cells[2].text = "%s" % rules
        elif (vuln == 'NEAA' or vuln == 'EEAA') and any_destination_any_port_table is not None:
            for acl in filter_rule_vulns[vuln]:
                row_cells = any_destination_any_port_table.add_row().cells
                row_cells[1].text = "%s" % acl
                rules = ""
                rule_device = None
                for rule in filter_rule_vulns[vuln][acl]:
                    if len(rules) > 0:
                        rules += ", %s" % rule
                    else:
                        rules += "%s" % rule
                    if rule_device is None:
                        if filter_rule_vulns[vuln][acl][rule]['Device']:
                            row_cells[0].text = "%s" % filter_rule_vulns[vuln][acl][rule]['Device']
                row_cells[2].text = "%s" % rules
        elif vuln == 'NENA' and any_port_table is not None:
            for acl in filter_rule_vulns[vuln]:
                row_cells = any_port_table.add_row().cells
                row_cells[1].text = "%s" % acl
                rules = ""
                rule_device = None
                for rule in filter_rule_vulns[vuln][acl]:
                    if len(rules) > 0:
                        rules += ", %s" % rule
                    else:
                        rules += "%s" % rule
                    if rule_device is None:
                        if filter_rule_vulns[vuln][acl][rule]['Device']:
                            row_cells[0].text = "%s" % filter_rule_vulns[vuln][acl][rule]['Device']
                row_cells[2].text = "%s" % rules
    return report


def write_vpn_aggressive(report, vulns):
    """Write Aggressive mode VPN section of the report"""

    report.add_heading("IKE Aggressive Mode Enabled", 3)

    report.add_paragraph("The assessment team identified %d utilizing the Internet key exchange (IKE) protocol "
                         "supporting aggressive mode with pre-shared key (PSK) authentication. Aggressive mode is a "
                         "feature that condenses the connection setup by immediately sending an unencrypted "
                         "authentication hash, which is derived using the PSK. In some instances, the PSK used to "
                         "establish a virtual private network (VPN) connection can be recovered by performing a "
                         "brute force attack on the authentication hash. Depending on the VPN gateway's configuration, "
                         "this could allow an attacker to gain unauthorized access to the network.", style='Normal')

    for vuln in vulns:
        pass
    return report


def write_weak_snmp(report, vulns):
    """Write the report verbiage for Weak SNMP Community String"""

    report.add_heading("SNMP Weak Community String", 3)

    report.add_paragraph("The assessment team identified %d utilizing a weak SNMP community string. SNMP community "
                         "strings are used like a password to restrict access to a host's management data. A brute "
                         "force attack can be used by an attacker to recover a community string, especially when it "
                         "is weak. After the community string is recovered it can be used by an attacker to enumerate "
                         "system information in support of other attacks, or even potentially change device "
                         "configurations to favorable conditions for the attacker.", style='Normal')

    report.add_heading("SNMP Weak Community String Affected Hosts", 4)
    table = report.add_table(rows=1, cols=3)
    table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Device'
    hdr_cells[1].text = 'Community String'
    hdr_cells[2].text = 'Weakness'

    for device in vulns:
        for community in vulns[device]:
            row_cells = table.add_row().cells
            row_cells[0].text = "%s" % device
            row_cells[1].text = "%s" % community
            row_cells[2].text = "%s" % vulns[device][community]['Weakness']

    report.add_paragraph("The assessment team recommends that all community strings for devices utilizing SNMP be "
                         "changed to follow industry standards and best practices for passwords.")

    return report


def save_report(report_object):
    """Save the generated assessment report"""
    out_dir = get_path()
    report_file = os.path.join(out_dir, "Nippy_Firewall_Configuraiton_Audit.docx")
    report_object.save(report_file)
    print warn + "Report saved to: " + report_file
    raw_input(info + "Press enter to continue...")


def get_path():
    """Prompt the user to enter a directory path"""

    output_path = None
    while output_path is None:
        print question + "Please enter the directory where you would like the file saved?"
        output_path = raw_input()
        if os.path.isdir(os.path.expanduser(output_path)):
            pass
        else:
            os.system('clear')
            print warn + "%s is not valid, please try again: " % str(output_path)
            output_path = None
    return os.path.expanduser(output_path)


if __name__ == '__main__':
    # Parse command line arguments
    parser = argparse.ArgumentParser()
    parser.add_argument('-X', '--xml', type=argparse.FileType('r'), required=True, help="Nipper XML Database file")
    parser.add_argument('-v', '--verbose', action='store_true', default=False, help="Enable verbose output to console")
    parser.add_argument('--debug', action='store_true', default=False, help="Enable debug output to console")
    args = parser.parse_args()

    try:
        report_data = {'Vulns': {}}

        logging.debug("Reading file: %s..." % args.xml.name)
        if args.debug:
            DEBUG = True
            print note + "Reading %s..." % args.xml.name
        parse_xml(args.xml)
        if args.verbose:
            VERBOSE = True
        if DEBUG:
            nv = get_nipper_version()
            print info + "Nipper Version: %s" % nv

        report_data['Devices'] = get_devices_summary()

        # report_data['Devices'] = get_devices_information() # Replaced with get_devices_summary function

        logging.info("%s device(s) found!" % str(len(report_data['Devices'])))
        if VERBOSE or DEBUG:
            print info + "%s device(s) found!" % str(len(report_data['Devices']))

        audit_sections = []
        for ref in nipper_xml.findall("./report/part/[@ref='SECURITYAUDIT']/section"):
            audit_sections.append(ref.get('ref'))
        if DEBUG:
            print info + "Audit Sections: %s" % audit_sections
            raw_input(warn + "Press enter to continue")

        for device in report_data['Devices'].keys():
            report_data['Devices'][device]['Interfaces'] = get_ipv4_interfaces(device)

        # Any Source, Any Destination, Any Port
        if "FILTER.RULE.AEAA" in audit_sections:
            report_data['Vulns']['AEAA'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.AEAA")
        # Any Source, Fixed Destination, Any Port
        if "FILTER.RULE.AENA" in audit_sections:
            report_data['Vulns']['AENA'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.AENA")
        # Fixed Source, Any Destination, Any Port
        if "FILTER.RULE.NEAA" in audit_sections:
            report_data['Vulns']['NEAA'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.NEAA")
        # Fixed Source, Fixed Destination, Any Port
        if "FILTER.RULE.NENA" in audit_sections:
            report_data['Vulns']['NENA'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.NENA")
        # Any Source, Any Destination
        if "FILTER.RULE.AEAE" in audit_sections:
            report_data['Vulns']['AEAE'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.AEAE")
        # Any Destination, Any Port
        if "FILTER.RULE.EEAA" in audit_sections:
            report_data['Vulns']['EEAA'] = get_filter_rules(report_data['Devices'].keys(), "FILTER.RULE.EEAA")
        # Aggressive mode VPN
        if "REMOTEACCESS.AGGRESSIVEMODE" in audit_sections:
            report_data['Vulns']['REMOTEACCESS.AGGRESSIVEMODE'] = get_vpn_aggressive(report_data['Devices'].keys())
        # Weak SNMP Community String
        if "SNMP.WEAK.COMMUNITY" in audit_sections:
            report_data['Vulns']['SNMP.WEAK.COMMUNITY'] = get_weak_snmp_community(report_data['Devices'].keys())
        report_data['CVE'] = get_cve(report_data['Devices'].keys())
        report_object = docx.Document()
        docx_file = write_intro(report_object, report_data['Devices'])
        docx_file = write_interfaces(report_object, report_data['Devices'])
        docx_file = write_filter_rules(report_object, report_data['Vulns'])
        if "SNMP.WEAK.COMMUNITY" in audit_sections:
            docx_file = write_weak_snmp(report_object, report_data['Vulns']['SNMP.WEAK.COMMUNITY'])
        # TODO Finish implementing the Aggressive Mode functionality
        # if report_data['Vulns']['REMOTEACCESS.AGGRESSIVEMODE'] :
        #    docx_file = write_vpn_aggressive(docx_file, report_data['Vulns']['REMOTEACCESS.AGGRESSIVEMODE'])
        if len(report_data['CVE']) > 0:
            docx_file = write_cve(report_object, report_data['CVE'])
        save_report(report_object)
        if DEBUG:
            print note + "Report_Data object: ", report_data
    except KeyboardInterrupt:
        print "\n" + warn + "User Interrupt! Quitting...."
    except:
        print "\n" + warn + "Please report this error to " + __maintainer__ + " by email at: " + __email__
        raise

#DATA STRUCTURE OF DEVICES
#report_data: {}
#   Devices: {}
    #   [<device name>]
    #       name: <name>
    #       type: <type>
    #       os: <os>
    #       osversion: <osversion>
    #       interfaces: {}
    #           [<interface>]
    #               interface: <interface>
    #               active: <active>
    #               address: <address>
    #               aclin: <acl in>
    #               aclout: <acl out>
#   CVE: {}
#       <CVE_ID>: {}
#           hosts: <hosts>
#           CVSSv2_Score: <score>
#   VULNS: {}
#       <VULN_ID>: {}
#           <ACL_ID>: {}
#               <Rule_ID>: {}
#                   <Rule Data>
#                   device: <device>
