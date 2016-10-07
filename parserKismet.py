#!/usr/bin/python
#!/usr/bin/env python

"""Parse Kismet netxml Files"""

from xml.etree import ElementTree
import argparse
import sys
import os
import hashlib

#################################################
#                    Variables                  #
#################################################
__author__ = "Russel Van Tuyl"
__license__ = "GPL"
__version__ = "0.1"
__maintainer__ = "Russel Van Tuyl"
__email__ = "Russel.VanTuyl@gmail.com"
__status__ = "Development"
DEBUG = False
#################################################
#                   COLORS                      #
#################################################
note = "\033[0;0;33m[-]\033[0m"
warn = "\033[0;0;31m[!]\033[0m"
info = "\033[0;0;36m[i]\033[0m"
question = "\033[0;0;37m[?]\033[0m"
debug = "\033[0;0;31m[DEBUG]\033[0m"


def parse_xml(file_object):
    """Parse an XML file and return the object"""
    if args.verbose:
        print note + 'Entered into parse_xml function'

    print note + "Parsing %s" % file_object.name

    try:
        xml_file = open(file_object.name, "r")
        xml = ElementTree.parse(xml_file)
        return xml
    except ElementTree.ParseError as e:
        print warn + "There was an error parsing the file: %s" % file_object
        raw_input("Press enter to continue")


def parse_directory(directory):
    """Parse all netxml Files in directory"""

    files = None
    k_files = []
    if os.path.isdir(os.path.expanduser(directory)):
        files = os.listdir(directory)

    if files is not None:
        for f in files:
            if f.lower().endswith('.netxml'):
                k_files.append(os.path.join(os.path.expanduser(directory), f))

    return k_files


def get_essid(xml_object):
    """Get a unique list of ESSIDs"""

    if args.verbose:
        print note + "Entered into get_essid function"
        print note + "%s" % xml_object

    for essid in xml_object.findall('.//wireless-network/SSID/essid'):  # /wireless-network/SSID/essid
        if args.verbose:
            print info + essid.text
        if essid.text not in kismet:
            kismet[essid.text] = {}


def get_data(xml_object):
    """Take in an XML object and transform it into a python data object"""

    if args.verbose:
        print note + "Entered into get_data function"
        print note + "%s" % xml_object

    kismet = {}

    for w in xml_object.findall(".//wireless-network/[@type='infrastructure']"):
        essid = None
        channel = None
        encryption = []
        bssid = None
        if w.find('./SSID/essid') is not None:
            essid = w.find('./SSID/essid').text
            if args.verbose:
                print info + "%s" % essid
        if w.find('./SSID/encryption') is not None:
            for e in w.findall('./SSID/encryption'):
                encryption.append(e.text)
                if args.verbose:
                    print "\t" + info + "Encryption: %s" % e.text
        if w.find('./BSSID') is not None:
            bssid = w.find('./BSSID').text
            if args.verbose:
                print "\t" + info + "BSSID: %s" % bssid
        if w.find('./channel') is not None:
            channel = w.find('./channel').text
            if args.verbose:
                print "\t" + info + "Channel: %s" % channel
        if essid and essid in kismet:
            if bssid not in kismet[essid] and bssid:
                kismet[essid][bssid] = {'channel': channel, 'encryption': encryption}
        elif essid:
            kismet[essid] = {}
            kismet[essid][bssid] = {'channel': channel, 'encryption': encryption}

    if DEBUG:
        for k in kismet:
            print debug + info + k
            for bssid in kismet[k]:
                print "\t %s: %s" % (bssid, kismet[k][bssid])
            raw_input(debug + "Press any key to continue...")

    return kismet


def print_to_screen(kismet):
    """Print the Kismet dictionary standard out"""

    print "\nESSID\t\tChannel(s)\t\tBSSID(s)"
    for essid in kismet:
        bssids = []
        channels = []
        for bssid in kismet[essid]:
            bssids.append(bssid)
            for channel in kismet[essid][bssid]['channel']:
                channels.append(channel)
        print "%s\t\t%s\t\t%s" % (essid, ', '.join(set(channels)), ', '.join(bssids))


def write_word_doc(kismet):
    """Write parsked Kismet data to a Word document"""

    import docx
    import math
    import readline

    readline.parse_and_bind('tab: complete')
    readline.set_completer_delims('\t')

    report = docx.Document()
    report.add_heading("Parsed Kismet Data", level=0)
    report.add_heading("Table of %s ESSIDs" % len(kismet), level=2)

    k = []

    for essid in kismet:
        if essid not in k:
            k.append(essid)

    c = 4  # number of desired columns
    r = int(math.ceil((len(kismet) / float(4))))  # Determine number of rows for table using a max of 4 columns
    table = report.add_table(rows=r, cols=c)
    table.style = 'Medium Grid 1 Accent 1'
    z = 0  # number of hosts
    x = 0  # row indices
    y = 0  # column indices
    while z < len(k):
        if (y / float(c)) == 1:  # Determine if we need to start putting data on a new row
            y = 0  # reset column indices since max number of columns reached
            x += 1
        table.cell(x, y).text = k[z]
        z += 1
        y += 1  # Add one to up the column data is put in
    if len(k) / float(c) != 1.000:  # Add "---" for empty spots in table
        d = c * (x + 1)
        while d > len(k):
            table.cell(x, y).text = "---"
            d -= 1
            y += 1

    # Print full table of BSSIDs
    report.add_heading("Table of All BSSIDs", level=2)
    t = report.add_table(rows=1, cols=4)
    t.style = 'Medium Grid 1 Accent 1'
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'ESSID'
    hdr_cells[1].text = 'BSSID'
    hdr_cells[2].text = 'Channel'
    hdr_cells[3].text = 'Security'
    table.style = 'Medium Grid 1 Accent 1'
    for essid in kismet:
        for bssid in kismet[essid]:
            row_cells = t.add_row().cells
            row_cells[0].text = "%s" % essid
            row_cells[1].text = "%s" % bssid
            row_cells[2].text = "%s" % kismet[essid][bssid]['channel']
            row_cells[3].text = "%s" % kismet[essid][bssid]['encryption']

    report.add_heading("Individual ESSID Tables", level=2)
    for essid in kismet:
        if args.essid:
            if essid in args.essid:
                if args.verbose:
                    print info + "Found ESSID: %s" % essid
                report.add_heading("Table for ESSID: %s" % essid, level=3)
                table = report.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ESSID'
                hdr_cells[1].text = 'BSSID'
                hdr_cells[2].text = 'Channel'
                hdr_cells[3].text = 'Security'
                table.style = 'Medium Grid 1 Accent 1'
                for bssid in kismet[essid]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = "%s" % essid
                    row_cells[1].text = "%s" % bssid
                    row_cells[2].text = "%s" % kismet[essid][bssid]['channel']
                    row_cells[3].text = "%s" % kismet[essid][bssid]['encryption']
        else:
            report.add_heading("Table for ESSID: %s" % essid, level=3)
            table = report.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ESSID'
            hdr_cells[1].text = 'BSSID'
            hdr_cells[2].text = 'Channel'
            hdr_cells[3].text = 'Security'
            table.style = 'Medium Grid 1 Accent 1'
            for bssid in kismet[essid]:
                row_cells = table.add_row().cells
                row_cells[0].text = "%s" % essid
                row_cells[1].text = "%s" % bssid
                row_cells[2].text = "%s" % kismet[essid][bssid]['channel']
                row_cells[3].text = "%s" % kismet[essid][bssid]['encryption']

    save_report(report)


def get_path():
    """Prompt the user to enter a directory path"""

    output_path = None
    while output_path is None:
        print question + "Please enter the directory where you would like the file saved?"
        output_path = raw_input()
        if os.path.isdir(os.path.expanduser(output_path)):
            pass
        else:
            os.system('clear')
            print warn + "%s is not valid, please try again: " % str(output_path)
            output_path = None
    return os.path.expanduser(output_path)


def save_report(report_object):
    """Save the generated report"""
    out_dir = get_path()
    report_file = os.path.join(out_dir, "Kismet_Parsed_Data.docx")
    report_object.save(report_file)
    print warn + "Report saved to: " + report_file
    raw_input(info + "Press enter to continue...")


def csv_list(string):
    return string.split(',')


if __name__ == '__main__':
    """Main Kismet netxml parser when run as a script"""

    parser = argparse.ArgumentParser()
    parser.add_argument('-X', '--xml', type=argparse.FileType('r'), help="Kismet .netxml file")
    parser.add_argument('-D', '--directory', help="Directory containing Kismet .netxml files")
    parser.add_argument('-E', '--essid', type=csv_list,
                        help="A comma seperated list of essids to filter output on")
    parser.add_argument('-W', '--word', action='store_true', default=False,
                        help="Save parsed results to a Word document")
    parser.add_argument('-C', '--csv', action='store_true', default=False,
                        help="Save parsed results to a CSV document on stdout")
    parser.add_argument('-v', '--verbose', action='store_true', default=False, help="Verbose Output")
    parser.add_argument('--debug', action='store_true', default=False, help="Enable debug output")
    args = parser.parse_args()

    try:
        DEBUG = args.debug

        if args.xml:
            netxml = parse_xml(args.xml)
            kismet = get_data(netxml)
            if args.word:
                write_word_doc(kismet)
            elif args.csv:
                print_to_screen(kismet)  # TODO replace with a real write csv to file function
        elif args.directory:
            kismet = {}
            kismet_files = parse_directory(args.directory)
            for f in kismet_files:
                netxml = parse_xml(open(f, "r"))
                k = get_data(netxml)
                for i in k:
                    if i not in kismet:
                        kismet[i] = k[i]
                    elif i in kismet:
                        kismet[i].update(k[i])
            if args.word:
                write_word_doc(kismet)
            elif args.csv:
                print_to_screen(kismet)
        else:
            print warn + "No arguments provided!"
            print warn + "Try: python " + __file__ + " --help"

    except KeyboardInterrupt:
        print warn + "\nUser Interrupt! Quitting...."
    except:
        print info + "\nPlease report this error to " + __maintainer__ + " by email at: " + __email__
        raise

# Data Structure
# <ESSID> : {}
#   <BISSID> : {}
#       <channel> : channel
#       <encryption> : encryption